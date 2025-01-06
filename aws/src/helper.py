import tika

tika.initVM()
from tika import parser
from openai_client import client
import pdfplumber
import json
from docx import Document
import csv
from datetime import datetime
import time

ASSISTANCE = "asst_GWO6ZchMGi1XGVoKgPXMLsOe"

"""
The tika_parser &  tika_bytes_parser function uses Apache Tika to extract and clean text
"""


def tika_parser(filename):
    text_from_document = None
    try:
        tika.initVM()
        headers = {"X-Tika-OCRLanguage": "eng", "X-Tika-PDFextractInlineImages": "true"}
        doc1 = parser.from_file(filename, headers=headers)
        while "\n\n" in doc1["content"]:
            doc1["content"] = doc1["content"].replace("\t", " ").replace("\n\n", "\n")
        while "  " in doc1["content"]:
            doc1["content"] = doc1["content"].replace("  ", " ")
        text_from_document = doc1["content"].encode().decode("ascii", "ignore")
        """ Text cleaning for mulitple spaces and . - . Add if anything else is needed"""
    except Exception as e:
        print("Error in tika", e)
        text_from_document = None
    return text_from_document


def tika_bytes_parser(bytes_io_obj):
    text_from_document = None
    try:
        tika.initVM()
        bytes_data = bytes_io_obj.getvalue()
        headers = {"X-Tika-OCRLanguage": "eng", "X-Tika-PDFextractInlineImages": "true"}
        doc = parser.from_buffer(bytes_data, headers=headers)
        cleaned_text = (
            doc["content"].replace("\t", " ").replace("\n\n", "\n").replace("  ", " ")
        )
        text_from_document = cleaned_text.encode().decode("ascii", "ignore")
    except Exception as e:
        print("Error in tika bytes:", e)
        text_from_document = None
    return text_from_document


def extract_text_from_pdf(pdf_path, filename):
    text_from_document = ""
    try:
        if filename.lower().endswith(".pdf"):
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text_from_document += page.extract_text()
        elif filename.lower().endswith(".txt"):
            with open(filename, "r", encoding="utf-8") as data:
                text_from_document = data.read()
        elif filename.lower().endswith(".docx") or filename.lower().endswith(".doc"):
            doc = Document(pdf_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_from_document = "\n".join(full_text)
        else:
            text_from_document = ""
    except Exception as e:
        print("Error extracting text from PDF:", e)
    return text_from_document


def parser_function(filename):
    text_from_document = ""
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(filename) as pdf:
                for page in pdf.pages:
                    text_from_document += page.extract_text()
        elif filename.endswith(".txt"):
            with open(filename, "r", encoding="utf-8") as data:
                text_from_document = data.read()
        elif filename.endswith(".docx") or filename.endswith(".doc"):
            doc = Document(filename)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_from_document = "\n".join(full_text)
        else:
            text_from_document = ""
    except Exception as e:
        print("Error extracting text from PDF:", e)
    return text_from_document


"""
The FineTuningModal function generates responses using a specified language model 
based on system and user messages, tracking token usage and handling errors.
"""


def FineTuningModal(system, user_message, modal):
    try:
        completion = client.chat.completions.create(
            model=modal,
            temperature=0,
            top_p=0,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user_message},
            ],
        )
        response = (completion.choices[0].message).content
        if "json" in response:
            response = response.replace("json", "")
        response = response.replace("```", "")
        token = completion.usage.total_tokens
    except Exception as e:
        print("Exception in Finetune", e)
        response = None
        token = 0
    return response, token


""" 
b'[{"skills":10}]' decode bytes to string and parse JSON string into a Python list
"""


def bytes_extract(bytes_data):
    if isinstance(bytes_data, bytes):
        json_string = bytes_data.decode("utf-8")
        try:
            data_list = json.loads(json_string)
        except:
            data_list = json_string
        return data_list
    return bytes_data


"""
CSV download for DCQC report 
"""


def csv_writer(data):
    if isinstance(data, list):
        # Get the headers from the keys of the first dictionary
        headers = data[0].keys()
        file_path = "output.csv"
        # Writing to CSV file
        with open(file_path, "w", newline="") as csvfile:
            writer = csv.writer(csvfile)
            serial_number = 1
            headers = ["S.No"] + list(headers)
            writer.writerow(headers)
            for ids, i in enumerate(data):
                row = data[ids].values()
                row = [serial_number] + list(row)
                writer.writerow(row)
                serial_number += 1
        return file_path


#: Convert the 2024-07-16T06:05:52Z into  2024-07-16 format
def date_extract(timestamp):
    datetime_object = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
    date_only = datetime_object.date()
    return date_only


CATEGORY_DATA = {
  "TMF": {
    "Global": {
      "01 Trial Management": {
        "01.01 Trial Management Plan": {},
        "01.02 Trial Master File Plan": {},
        "01.03 Trial Oversight Committee Documentation": {},
        "01.04 Trial Team Training Records": {},
        "01.05 Trial Team Meeting Minutes": {},
        "01.06 Trial Communication Plan": {},
        "01.07 Trial Risk Management Plan": {},
        "01.08 Trial Monitoring Plan": {},
        "01.09 Trial Monitoring Reports": {},
        "01.10 Trial Monitoring Visit Log": {},
        "01.11 Trial Monitoring Issue Log": {},
        "01.12 Trial Monitoring Issue Resolution Documentation": {},
        "01.13 Trial Monitoring Close-Out Report": {},
        "01.14 Trial Close-Out Plan": {},
        "01.15 Trial Close-Out Report": {}
      },
      "02 Regulatory": {
        "02.01 Regulatory Authority Submission": {},
        "02.02 Regulatory Authority Approval": {},
        "02.03 Regulatory Authority Correspondence": {},
        "02.04 Regulatory Authority Inspection Documentation": {},
        "02.05 Regulatory Authority Inspection Response": {},
        "02.06 Regulatory Authority Inspection Close-Out Documentation": {}
      },
      "03 Ethics": {
        "03.01 Ethics Committee Submission": {},
        "03.02 Ethics Committee Approval": {},
        "03.03 Ethics Committee Correspondence": {},
        "03.04 Ethics Committee Membership List": {},
        "03.05 Ethics Committee Meeting Minutes": {},
        "03.06 Ethics Committee Inspection Documentation": {},
        "03.07 Ethics Committee Inspection Response": {},
        "03.08 Ethics Committee Inspection Close-Out Documentation": {}
      },
      "04 Trial Design": {
        "04.01 Protocol": {},
        "04.02 Protocol Amendments": {},
        "04.03 Investigator’s Brochure": {},
        "04.04 Investigational Medicinal Product Dossier": {},
        "04.05 Investigational Device Dossier": {},
        "04.06 Sample Case Report Form": {},
        "04.07 Sample Informed Consent Form": {},
        "04.08 Sample Patient Information Sheet": {},
        "04.09 Sample Recruitment Materials": {},
        "04.10 Sample Data Collection Tools": {},
        "04.11 Sample Patient Diary": {},
        "04.12 Sample Patient Reported Outcome Measures": {}
      },
      "05 Safety Reporting": {
        "05.01 Safety Management Plan": {},
        "05.02 Adverse Event Reporting Forms": {},
        "05.03 Serious Adverse Event Reporting Forms": {},
        "05.04 Suspected Unexpected Serious Adverse Reaction Reports": {},
        "05.05 Data Safety Monitoring Board Charter": {},
        "05.06 Data Safety Monitoring Board Meeting Minutes": {},
        "05.07 Data Safety Monitoring Board Reports": {},
        "05.08 Safety Reporting Correspondence": {},
        "05.09 Safety Reporting to Regulatory Authorities": {},
        "05.10 Safety Reporting to Ethics Committees": {},
        "05.11 Safety Reporting to Investigators": {},
        "05.12 Safety Reporting to Data Safety Monitoring Board": {}
      },
      "06 Investigational Product": {
        "06.01 Investigational Product Management Plan": {},
        "06.02 Investigational Product Shipping Records": {},
        "06.03 Investigational Product Accountability Logs": {},
        "06.04 Investigational Product Storage Conditions Documentation": {},
        "06.05 Investigational Product Temperature Deviation Reports": {},
        "06.06 Investigational Product Destruction Records": {},
        "06.07 Investigational Product Return Records": {},
        "06.08 Investigational Product Recall Documentation": {},
        "06.09 Investigational Product Labeling Documentation": {},
        "06.10 Investigational Product Randomization Codes": {},
        "06.11 Investigational Product Blinding and Unblinding Documentation": {}
      },
      "07 Laboratory": {
        "07.01 Central Laboratory Agreements": {},
        "07.02 Central Laboratory Accreditation Certificates": {},
        "07.03 Central Laboratory Normal Value Ranges": {},
        "07.04 Central Laboratory Instructions": {},
        "07.05 Central Laboratory Sample Shipment Records": {},
        "07.06 Central Laboratory Sample Analysis Reports": {},
        "07.07 Central Laboratory Quality Control Documentation": {},
        "07.08 Central Laboratory Correspondence": {}
      },
      "08 Data Management": {
        "08.01 Data Management Plan": {},
        "08.02 Data Management System Validation Documentation": {},
        "08.03 Data Entry Guidelines": {},
        "08.04 Data Query Forms": {},
        "08.05 Data Query Resolution Documentation": {},
        "08.06 Data Transfer Agreements": {},
        "08.07 Data Transfer Records": {},
        "08.08 Data Lock Documentation": {},
        "08.09 Data Archiving Plan": {},
        "08.10 Data Archiving Confirmation": {}
      },
      "09 Statistics": {
        "09.01 Statistical Analysis Plan": {},
        "09.02 Randomization List": {},
        "09.03 Blinding and Unblinding Procedures": {},
        "09.04 Interim Statistical Analysis Reports": {},
        "09.05 Final Statistical Analysis Report": {},
        "09.06 Statistical Programming Validation Documentation": {}
      },
      "10 Clinical Study Report": {
        "10.01 Clinical Study Report": {},
        "10.02 Clinical Study Report Appendices": {},
        "10.03 Clinical Study Report Submission Documentation": {},
        "10.04 Clinical Study Report Approval Documentation": {}
      },
      "11 Site Management": {
        "11.01 Site Selection Plan": {},
        "11.02 Site Feasibility Questionnaires": {},
        "11.03 Site Selection Reports": {},
        "11.04 Site Initiation Visit Reports": {},
        "11.05 Site Training Records": {},
        "11.06 Site Monitoring Visit Reports": {},
        "11.07 Site Close-Out Visit Reports": {},
        "11.08 Site Correspondence": {},
        "11.09 Site Regulatory Documents": {},
        "11.10 Site Informed Consent Forms": {},
        "11.11 Site Source Documents": {},
        "11.12 Site Screening and Enrollment Logs": {},
        "11.13 Site Delegation of Authority Logs": {},
        "11.14 Site Signature Logs": {},
        "11.15 Site Financial Agreements": {},
        "11.16 Site Payment Records": {},
        "11.17 Site Audit Reports": {},
        "11.18 Site Audit Response Documentation": {}
      },
      "12 Quality Assurance": {
        "12.01 Quality Management Plan": {},
        "12.02 Standard Operating Procedures": {},
        "12.03 Training Records": {},
        "12.04 Audit Plans": {},
        "12.05 Audit Reports": {},
        "12.06 Audit Response Documentation": {},
        "12.07 Inspection Plans": {},
        "12.08 Inspection Reports": {}
      }
    },
    "Region_1": "(Replicate entire Global folder structure here)",
    "Region_2": "(Replicate entire Global folder structure here)",
    "Region_3": "(Replicate entire Global folder structure here)"
  }
}







#TODO: Due to Incomplete Modal change a manual Datasets
def DCQC_Modal(message,index):
    is_active  = False
    if is_active: ## For Clinincal Prompt
        system = 'You are a highly specialized TMF and QMS document classifier'
        user_message = '''
        You are a highly specialized TMF and QMS document classifier built for the biotech and pharmaceutical industries. When a document is uploaded, you must perform the following tasks with precision:

        1. Identify and Classify Documents:
        Determine the document type (e.g., SOP, Clinical Study Report, Investigator Brochure) and classify it accurately according to the TMF or QMS reference models.

        2. Recommend Archival Locations:
        Suggest the appropriate folder and section numbers (e.g., TMF Section 8.2.1 or QMS 4.3.2) where the document should be stored. Ensure recommendations align with industry standards.

        3. Review for Compliance:
        Assess the document for ALCOA principles compliance and adherence to Good Document Management Practices.
        Highlight any deviations, such as illegibility, lack of timestamps, or missing critical information.

        4. Flag Drafts or Editable Documents:
        Clearly warn the user if the document appears to be a draft or is in an editable format. Use red warnings to emphasize the issue and explain how to convert it into a final version suitable for submission.

        5. Provide Tailored Feedback:
        Offer specific and actionable recommendations to improve the document's quality.
        Include suggestions to enhance clarity, formatting, regulatory compliance, and completeness.

        6. Engage Constructively:
        Use a professional and supportive tone to guide users, ensuring feedback is clear, actionable, and focused on improving document quality.

        Your goal is to ensure every document is submission-ready, compliant with industry standards, and properly archived for retrieval during audits or regulatory inspections."
        '''    


        output_format = '''
            {
                "type": "SOP",
                "classification": {
                "category": "TMF",
                "section": "01 Trial Management",
                "subsection": "01.02 Trial Master File Plan",
                "folder_path": "TMF/Global/01 Trial Management/01.02 Trial Master File Plan/"
                },
                "compliance_check": {
                "alcoa_principles": {
                    "attributable": true,
                    "legible": false,
                    "contemporaneous": true,
                    "original": true,
                    "accurate": true
                },
                "gdmp_adherence": {
                    "timestamp_presence": false,
                    "signature_presence": true,
                    "critical_information": false
                },
                "issues": [
                    "Illegibility detected.",
                    "Missing timestamps.",
                    "Critical information not found."
                ]
                },
                "status": {
                "is_draft": true,
                "warnings": [
                    {
                    "type": "Editable Format",
                    "message": "The document appears to be in an editable format. Convert it to PDF/A for final submission."
                    }
                ]
                },
                "feedback": {
                "quality_improvements": [
                    "Ensure all text is legible and formatted clearly.",
                    "Add timestamps to comply with ALCOA principles.",
                    "Include missing critical information such as approval date."
                ],
                "regulatory_compliance": [
                    "Verify adherence to Good Documentation Practices.",
                    "Review for completeness against the TMF or QMS model."
                ]
                }
            }
        '''
        user_message = user_message + f'\n Category of this  TMF and QMS. \n{CATEGORY_DATA}' +'\n You May Classify and categories  this section of the content and return the Output as Json Format which You Categories ' + output_format
    else:    ## For General Prompt
        system = 'You Are a Document Classifier Give a Name of the File'
        user_message = '''
        You are an intelligent categorization system. Your job is to analyze the given  file and categorize its content into one of the following categories:
        Resume: The content appears to be a candidate's resume, containing personal details, skills, education, work experience, certifications, etc.
        Job Description: The content describes a job role, responsibilities, qualifications, skills required, and related details.
        Other: The content does not clearly fit into the above two categories.
        Input:You will receive a  file with structured or unstructured content

        Output:Your output should be a single line indicating the category the  file belongs to, formatted as: 
        {"Category" : "Resume"}
        '''    
    final_text = user_message +  '\n' + message
    modal = "gpt-3.5-turbo"
    data,token = FineTuningModal(system,final_text,modal)
    if isinstance(data,str):
        data = json.loads(data)
    return data,token
        
    
def date_checking(date):
    today = datetime.now().date()
    try:
        date_from_data = datetime.strptime(date, "%Y-%m-%dT%H:%M:%SZ").date()
    except:
        date_from_data = datetime.strptime(date, "%Y-%m-%d %H:%M:%S").date()
    # Compare dates
    if date_from_data == today:
        return True
    else:
        return False


def ChatBotGpt(parse_text):
    user_message = parse_text
    try:
        # Create a new thread
        thread = client.beta.threads.create()
        thread_id = thread.id
 
        # Add user message to the thread
        messages_id = client.beta.threads.messages.create(
            thread_id=thread_id, content=user_message, role="user"
        )
 
        # Start a new run
        my_run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=ASSISTANCE,
        )

        # Check run status periodically
        while my_run.status in ["queued", "in_progress"]:
            keep_retrieving_run = client.beta.threads.runs.retrieve(
                thread_id=thread_id, run_id=my_run.id
            )

            if keep_retrieving_run.status == "completed":
                all_messages = client.beta.threads.messages.list(thread_id=thread_id)
                length = len(all_messages.data)
                for ids, i in enumerate(reversed(all_messages.data)):
                    if ids == length - 1:
                        response = i.content[0].text.value
                        break
                if response:
                    break
            time.sleep(5)

    except Exception as e:
        print("ChatBotGpt Exception:", e)
        return None

    return response
